from docx import Document
from pathlib import Path

# 1. 输入 Word 文件路径
input_file = Path("南通R18032地块项目与轨道交通1号线盾构区间相互影响安全评估报告20190624.docx")

# 2. 输出 txt 文件路径
output_file = Path("南通R18032安全评估报告_纯文本.txt")

doc = Document(input_file)

lines = []

# 3. 提取普通段落
for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if text:
        lines.append(text)

# 4. 提取表格文字
for table in doc.tables:
    lines.append("\n[表格开始]")
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            cell_text = cell.text.strip().replace("\n", " ")
            row_text.append(cell_text)
        lines.append(" | ".join(row_text))
    lines.append("[表格结束]\n")

# 5. 保存为 txt
output_file.write_text("\n".join(lines), encoding="utf-8")

print("转换完成：", output_file)