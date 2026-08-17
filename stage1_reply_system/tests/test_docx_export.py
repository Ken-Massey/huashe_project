import tempfile
import unittest
from pathlib import Path

from docx import Document

from stage1_reply_system.review_generation.docx_export import (
    render_audit_record_docx,
    render_reply_draft_docx,
)


def package() -> dict:
    return {
        "case_id": "DOCX-001",
        "overall_status": "ready_for_human_review",
        "formal_issuance_note": "系统输出为辅助审核草稿，须经专业人员复核。",
        "project_facts": [
            {"label": "项目名称", "value": "测试基坑工程"},
            {"label": "报审单位", "value": "测试建设单位"},
            {"label": "项目阶段", "value": "规划"},
        ],
        "audit_opinions": [{
            "topic": "退让距离/净距控制",
            "result": "满足",
            "conclusion": "审核结论：满足。",
            "regulation_clauses": ["3.3.1"],
            "function": "evaluate_setback_distance",
            "calculation_steps": ["8m 大于 6m"],
        }],
        "missing_required_inputs": [],
        "historical_advice": {
            "attention_items": ["（1）施工前开展现状调查。"],
            "source_project": "历史测试工程",
            "source_similarity": 0.9,
            "source_reply_file": "reply.docx",
            "extraction_anchor": "请注意如下事项",
        },
        "history_match": {
            "selected_match": {
                "history_case_id": 1,
                "history_quality_status": "ready",
                "history_quality_issues": [],
                "component_scores": {"project_stage": 1.0},
            }
        },
        "formal_reply": {
            "title": "关于测试基坑工程规划设计方案征求地铁意见的复函",
            "recipient": "测试建设单位",
            "introduction": "贵司关于测试基坑工程规划设计方案的函件及相关资料已收悉。经研究，具体意见函复如下：",
            "project_description": "项目位于测试路，邻近地铁1号线盾构区间。",
            "attention_lead": "我部原则同意该方案。为确保地铁结构及运营安全，请贵司在后续工作中注意如下事项：",
            "attention_items": ["施工前开展现状调查。", "施工期间开展变形监测。"],
            "closing_requirement": "该项目施工图设计、施工方案应书面征求我部意见。",
            "closing": "特此函复。",
            "attachment": "测试基坑工程与地铁平面关系图",
            "issuing_organization": "南京市地铁交通设施保护办公室",
            "issue_date": "2026年7月24日",
        },
    }


class DocxExportTests(unittest.TestCase):
    def test_generates_editable_reply_and_internal_record(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            reply = render_reply_draft_docx(package(), root / "reply.docx")
            audit = render_audit_record_docx(package(), root / "audit.docx")
            self.assertTrue(reply.exists())
            self.assertTrue(audit.exists())
            reply_text = "\n".join(p.text for p in Document(reply).paragraphs)
            audit_text = "\n".join(p.text for p in Document(audit).paragraphs)
            self.assertIn("关于测试基坑工程规划设计方案征求地铁意见的复函", reply_text)
            self.assertIn("测试建设单位：", reply_text)
            self.assertIn("（1）施工前开展现状调查。", reply_text)
            self.assertIn("南京市地铁交通设施保护办公室", reply_text)
            self.assertEqual(len(Document(reply).tables), 0)
            self.assertIn("第一阶段自动审核记录", audit_text)
            self.assertIn("evaluate_setback_distance", "\n".join(
                cell.text for table in Document(audit).tables for row in table.rows for cell in row.cells
            ))

    def test_omits_attachment_line_when_no_real_attachment_exists(self):
        with tempfile.TemporaryDirectory() as directory:
            data = package()
            data["formal_reply"]["attachment"] = None
            reply = render_reply_draft_docx(data, Path(directory) / "reply.docx")
            reply_text = "\n".join(p.text for p in Document(reply).paragraphs)
            self.assertNotIn("附件：", reply_text)
            self.assertNotIn("与地铁平面关系图", reply_text)


if __name__ == "__main__":
    unittest.main()
