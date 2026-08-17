"""Extract text from historical Word and PDF reply files."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any


def _clean_text(text: str) -> str:
    text = text.replace("\r\x07", "\n").replace("\x07", "\n").replace("\r", "\n")
    text = re.sub(r"[\t\u00a0]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    blocks: list[str] = []
    blocks.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append("\t".join(cells))
    return _clean_text("\n".join(blocks))


def extract_doc_with_word(path: Path) -> str:
    """Read legacy .doc through locally installed Microsoft Word."""
    try:
        import win32com.client
    except ImportError as exc:
        raise RuntimeError("读取 .doc 需要在 Windows 安装 pywin32，并安装 Microsoft Word。") from exc

    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(str(path.resolve()), False, True)
        return _clean_text(document.Content.Text)
    finally:
        if document is not None:
            document.Close(False)
        if word is not None:
            word.Quit()


def extract_pdf_text(path: Path, allow_ocr: bool = True) -> tuple[str, str]:
    from stage1_reply_system.document_processing.pdf_extractor import extract_pdf

    # Historical replies are short. Lower-resolution local OCR and a page cap
    # keep bulk indexing practical; the normal letter pipeline still uses 300 DPI.
    result = extract_pdf(
        path,
        direct_text_threshold=30 if allow_ocr else 0,
        ocr_resolution=180,
        max_pages=8,
        scan_last_page_footer=False,
    )
    pages = result.get("pages", [])
    text = "\n\n".join(page.get("text", "") for page in pages if page.get("text"))
    methods = {page.get("method") for page in pages if page.get("method")}
    method = methods.pop() if len(methods) == 1 else "mixed"
    return _clean_text(text), method or "direct"


def extract_document(path: str | Path, allow_pdf_ocr: bool = True) -> dict[str, Any]:
    source = Path(path)
    suffix = source.suffix.lower()
    if suffix == ".docx":
        text = extract_docx(source)
        method = "docx"
    elif suffix == ".doc":
        text = extract_doc_with_word(source)
        method = "word_com"
    elif suffix == ".pdf":
        text, method = extract_pdf_text(source, allow_ocr=allow_pdf_ocr)
        method = f"pdf_{method}"
    else:
        raise ValueError(f"不支持的历史文档格式：{source.suffix}")
    return {"path": str(source.resolve()), "text": text, "method": method}
