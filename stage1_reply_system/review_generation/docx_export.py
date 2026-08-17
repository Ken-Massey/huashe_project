"""Generate editable Word drafts and traceable internal audit records."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


STATUS_LABELS = {
    "blocked_by_missing_data": "关键资料待补充",
    "requires_special_study": "需专题专项论证",
    "ready_for_human_review": "计算完成，待人工复核",
}

SCORE_LABELS = {
    "project_name": "项目名称",
    "project_stage": "项目阶段",
    "project_type": "项目类型",
    "relative_relationship": "相对关系",
    "structure_method": "结构形式",
    "metro_line": "地铁线路",
    "pit_depth": "基坑深度",
    "horizontal_clearance": "水平净距",
    "vertical_clearance": "竖向净距",
    "letter_semantics": "函件语义",
}


def _set_run_font(run, name: str = "宋体", size: float = 11, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths_dxa)))
    width.set(qn("w:type"), "dxa")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Cm(widths_dxa[index] / 1440 * 2.54)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell._tc.get_or_add_tcPr().get_or_add_tcW().set(qn("w:w"), str(widths_dxa[index]))
            _set_cell_margins(cell)


def _keep_table_together(table) -> None:
    """Move a compact decision table as a unit when it cannot fit the page."""
    paragraphs = [paragraph for row in table.rows for cell in row.cells for paragraph in cell.paragraphs]
    for paragraph in paragraphs[:-1]:
        paragraph.paragraph_format.keep_with_next = True


def _configure_document(document: Document, *, internal: bool) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5" if internal else "000000", 16, 8),
        ("Heading 2", 13, "2E74B5" if internal else "000000", 12, 6),
        ("Heading 3", 12, "1F4D78" if internal else "000000", 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if internal:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run("内部审核记录  |  系统辅助生成")
        _set_run_font(run, "微软雅黑", 8)
        run.font.color.rgb = RGBColor(100, 106, 115)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if internal:
        run = footer.add_run("本文件须经专业人员复核、审批后使用")
        _set_run_font(run, "宋体", 8)
        run.font.color.rgb = RGBColor(110, 110, 110)
    else:
        _add_page_number(footer)


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run("— ")
    _set_run_font(run, "宋体", 10.5)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    run = paragraph.add_run(" —")
    _set_run_font(run, "宋体", 10.5)


def _add_bottom_border(paragraph, color: str = "D94A38", size: str = "12") -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)


def _set_formal_paragraph(paragraph, *, indent: bool = True, hanging: bool = False) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(28)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if hanging:
        paragraph.paragraph_format.left_indent = Pt(32)
        paragraph.paragraph_format.first_line_indent = Pt(-32)
    elif indent:
        paragraph.paragraph_format.first_line_indent = Pt(32)
    for run in paragraph.runs:
        _set_run_font(run, "仿宋_GB2312", 16)


def _add_formal_paragraph(
    document: Document,
    text: str,
    *,
    indent: bool = True,
    hanging: bool = False,
):
    paragraph = document.add_paragraph(str(text))
    _set_formal_paragraph(paragraph, indent=indent, hanging=hanging)
    return paragraph


def _without_list_marker(value: object) -> str:
    return re.sub(r"^\s*(?:[（(]\d+[）)]|\d+[.、])\s*", "", str(value)).strip()


def _add_reply_heading(document: Document, reply: dict[str, Any]) -> None:
    issuer_name = reply.get("issuing_organization") or "南京市地下铁道工程建设指挥部"
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(20)
    heading.paragraph_format.space_after = Pt(20)
    run = heading.add_run(str(issuer_name))
    _set_run_font(run, "方正小标宋简体", 28, False)
    run.font.color.rgb = RGBColor.from_string("D94A38")

    number = document.add_paragraph()
    number.alignment = WD_ALIGN_PARAGRAPH.CENTER
    number.paragraph_format.space_after = Pt(8)
    run = number.add_run(str(reply.get("document_number") or "宁地铁保护〔    〕    号"))
    _set_run_font(run, "仿宋_GB2312", 14, False)

    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(16)
    _add_bottom_border(rule)


def _add_title(document: Document, text: str, *, internal: bool) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(16)
    run = paragraph.add_run(text)
    _set_run_font(run, "黑体", 18 if internal else 20, True)
    if internal:
        run.font.color.rgb = RGBColor.from_string("1F4D78")


def _add_warning(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    _set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    _shade_cell(cell, "FFF2CC")
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    _set_run_font(run, "微软雅黑", 10, True)
    run.font.color.rgb = RGBColor.from_string("7F6000")


def _add_key_value_table(
    document: Document,
    rows: list[tuple[str, object]],
    *,
    keep_together: bool = False,
):
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = "" if value is None else str(value)
        _shade_cell(cells[0], "F2F4F7")
        for run in cells[0].paragraphs[0].runs:
            _set_run_font(run, "微软雅黑", 9.5, True)
        for run in cells[1].paragraphs[0].runs:
            _set_run_font(run, "宋体", 10)
    _set_table_geometry(table, [2160, 7200])
    if keep_together:
        _keep_table_together(table)
    return table


def render_reply_draft_docx(package: dict[str, Any], output_file: str | Path) -> Path:
    """Create an editable formal reply draft in the outgoing-letter layout."""
    path = Path(output_file)
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _configure_document(document, internal=False)
    facts = {fact["label"]: fact["value"] for fact in package.get("project_facts", [])}
    name = facts.get("项目名称", "项目名称待补充")
    stage = facts.get("项目阶段", "规划")
    applicant = facts.get("报审单位", "收函单位待补充")
    reply = package.get("formal_reply") or {}

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    title_paragraph.paragraph_format.line_spacing = Pt(30)
    title_paragraph.paragraph_format.space_after = Pt(22)
    run = title_paragraph.add_run(
        str(reply.get("title") or f"关于{name}{stage}方案征求地铁意见的复函")
    )
    _set_run_font(run, "方正小标宋简体", 20, False)

    _add_formal_paragraph(
        document,
        f"{reply.get('recipient') or applicant}：",
        indent=False,
    )
    _add_formal_paragraph(
        document,
        reply.get("introduction")
        or f"贵司关于{name}{stage}方案的函件及相关资料已收悉。经研究，具体意见函复如下：",
    )
    _add_formal_paragraph(
        document,
        f"1. {reply.get('project_description') or '项目建设内容及其与既有地铁结构的空间关系详见报审资料。'}",
    )
    _add_formal_paragraph(
        document,
        f"2. {reply.get('attention_lead') or '为确保地铁结构及运营安全，请贵司在后续工作中注意如下事项：'}",
    )
    attention_items = (
        reply.get("attention_items")
        or package.get("historical_advice", {}).get("attention_items")
        or []
    )
    for index, item in enumerate(attention_items, 1):
        _add_formal_paragraph(document, f"（{index}）{_without_list_marker(item)}", hanging=True)
    _add_formal_paragraph(
        document,
        f"3. {reply.get('closing_requirement') or ''}",
    )
    _add_formal_paragraph(document, reply.get("closing") or "特此函复。")

    attachment = str(reply.get("attachment") or "").strip()
    if attachment:
        document.add_paragraph()
        _add_formal_paragraph(document, f"附件：{attachment}")
    document.add_paragraph()
    issuer = _add_formal_paragraph(
        document,
        reply.get("issuing_organization") or "南京市地下铁道工程建设指挥部",
        indent=False,
    )
    issuer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    issue_date = _add_formal_paragraph(
        document,
        reply.get("issue_date") or "",
        indent=False,
    )
    issue_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.save(path)
    return path


def render_audit_record_docx(package: dict[str, Any], output_file: str | Path) -> Path:
    """Create an internal record containing decisions, functions and provenance."""
    path = Path(output_file)
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _configure_document(document, internal=True)
    _add_title(document, "第一阶段自动审核记录", internal=True)
    _add_warning(document, package["formal_issuance_note"])

    document.add_heading("1. 项目与审核状态", level=1)
    rows = [("案例编号", package["case_id"]), ("总体状态", STATUS_LABELS.get(package["overall_status"], package["overall_status"]))]
    rows.extend((fact["label"], fact["value"]) for fact in package["project_facts"])
    _add_key_value_table(document, rows)

    document.add_page_break()
    document.add_heading("2. 规程审核判断", level=1)
    for index, opinion in enumerate(package["audit_opinions"], 1):
        document.add_heading(f"2.{index} {opinion['topic']}", level=2)
        _add_key_value_table(document, [
            ("结论", opinion["result"]),
            ("审核意见", opinion["conclusion"]),
            ("依据条文", "、".join(opinion["regulation_clauses"])),
            ("审核函数", opinion["function"]),
            ("计算过程", "；".join(opinion["calculation_steps"]) or "无"),
        ], keep_together=True)

    document.add_heading("3. 缺失资料", level=1)
    if package["missing_required_inputs"]:
        for item in package["missing_required_inputs"]:
            document.add_paragraph(f"{item['label']}（{item['field']}）", style="List Bullet")
    else:
        document.add_paragraph("确定性计算所需字段已提供。")

    history = package["historical_advice"]
    match = package["history_match"].get("selected_match") or {}
    document.add_heading("4. 历史回函匹配与来源", level=1)
    _add_key_value_table(document, [
        ("历史案例编号", match.get("history_case_id")),
        ("来源项目", history.get("source_project")),
        ("总相似度", history.get("source_similarity")),
        ("质量状态", match.get("history_quality_status") or "未知（旧审核包未记录）"),
        ("质量提示", "、".join(match.get("history_quality_issues") or []) or "无"),
        ("来源回函", history.get("source_reply_file")),
        ("截取锚点", history.get("extraction_anchor")),
    ])
    scores = match.get("component_scores") or {}
    if scores:
        document.add_heading("4.1 匹配分项", level=2)
        _add_key_value_table(document, [
            (SCORE_LABELS.get(name, name), "无数据" if score is None else score)
            for name, score in scores.items()
        ])
    document.add_heading("4.2 历史注意事项原文", level=2)
    if history["attention_items"]:
        for item in history["attention_items"]:
            document.add_paragraph(str(item))
    else:
        document.add_paragraph("未自动采用历史注意事项。")

    document.add_page_break()
    document.add_heading("5. 人工复核", level=1)
    review_table = _add_key_value_table(document, [
        ("复核人", "待填写"),
        ("复核结论", "待填写"),
        ("修改说明", "待填写"),
        ("复核日期", "待填写"),
    ], keep_together=True)
    for row in review_table.rows:
        row.height = Cm(1.5)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    document.save(path)
    return path
