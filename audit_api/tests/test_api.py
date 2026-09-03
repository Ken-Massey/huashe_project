from __future__ import annotations

import json
import tempfile
import time
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from audit_api.agent import AgentService
from audit_api.agent_conversation import AgentConversationRepository
from audit_api.dynamic_audit import run_dynamic_regulation_audit
from audit_api.knowledge_base import KnowledgeBase
from audit_api.regulation_rules import RegulationRepository, RuleEngine, SafeExpression, extract_regulation
from audit_api.main import _artifact_files, app
from audit_api.services import _classify_project_document, _infer_document_type, _prepare_stage1_input, recognize_letter
from audit_api.task_manager import TaskManager
from deepke_case_extract.pipelines import run_audit_one_plan, run_match_new_case_advice


class ApiDefinitionTests(unittest.TestCase):
    def test_project_document_classifier_distinguishes_letter_and_case(self):
        letter_role, _, _ = _classify_project_document(
            "关于规划方案征求地铁意见的函.pdf",
            "现就该项目规划方案征求贵单位意见，请予复函。",
        )
        case_role, _, _ = _classify_project_document(
            "基坑工程安全影响评价报告.pdf",
            "本报告开展安全影响评价，采用有限元计算并提出监测方案。",
        )

        self.assertEqual(letter_role, "letter")
        self.assertEqual(case_role, "case")

    def test_document_type_prefers_filename_keywords(self):
        document_type, confidence, reason = _infer_document_type(
            "滨湖社区综合服务中心项目基坑工程涉地铁2号线结构安全性影响预评估报告.pdf",
            "工程概况 本项目基坑邻近地铁区间。",
        )
        self.assertEqual(document_type, "safety_assessment_report")
        self.assertGreaterEqual(confidence, 0.8)
        self.assertIn("文件名", reason)

        construction_type, _, _ = _infer_document_type(
            "滨湖社区卫生服务中心及基层社区中心项目桩基和支护施工方案.pdf",
            "本方案明确施工部署、施工工序及应急措施。",
        )
        self.assertEqual(construction_type, "construction_scheme")

    def test_recognize_letter_extracts_engineering_fields_from_plain_text(self):
        with tempfile.TemporaryDirectory() as folder:
            source = Path(folder) / "滨湖社区卫生服务中心项目基坑安全性影响预评估报告.txt"
            source.write_text(
                "工程概况\n"
                "本项目基坑位于地铁2号线云锦路站~莫愁湖站区间隧道北侧，"
                "基坑开挖深度约为6.7m，采用钻孔灌注桩、三轴搅拌桩止水帷幕，"
                "支护结构与地铁结构边线最小水平距离约为20.1m，降水方式为管井降水，"
                "对应地铁结构埋深约为14.5m。",
                encoding="utf-8",
            )

            result = recognize_letter(source)

        self.assertEqual(result["document_type"], "safety_assessment_report")
        self.assertEqual(result["document_role"], "case")
        fields = result["fields"]
        self.assertEqual(fields["project_type"], "基坑")
        self.assertEqual(fields["metro_line_name"], "2号线")
        self.assertIn("云锦路站~莫愁湖站", fields["metro_section_name"])
        self.assertEqual(fields["pit_depth_m"], 6.7)
        self.assertEqual(fields["minimum_horizontal_clearance_m"], 20.1)
        self.assertEqual(fields["dewatering_method"], "管井降水")
        self.assertIn("钻孔灌注桩", fields["support_components"])

    def test_explicit_artifact_files_hide_internal_task_outputs(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            opinion = root / "审核意见.md"
            reply = root / "最终版复函.docx"
            internal = root / "04_merged_input.json"
            opinion.write_text("审核意见", encoding="utf-8")
            reply.write_bytes(b"docx")
            internal.write_text("{}", encoding="utf-8")
            task = {
                "result": {
                    "artifact_files": [str(opinion), str(reply)],
                    "artifact_roots": [str(root)],
                }
            }

            files = _artifact_files(task)

            self.assertEqual([item["name"] for item in files], ["审核意见.md", "最终版复函.docx"])

    def test_legacy_stage1_task_also_only_exposes_two_public_files(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            (root / "审核意见.md").write_text("审核意见", encoding="utf-8")
            (root / "回函辅助草稿.docx").write_bytes(b"docx")
            (root / "00_manual_input.json").write_text("{}", encoding="utf-8")
            task = {
                "result": {
                    "stage": "stage1",
                    "artifact_roots": [str(root)],
                }
            }

            files = _artifact_files(task)

            self.assertEqual([item["name"] for item in files], ["审核意见.md", "回函辅助草稿.docx"])

    def test_required_routes_exist(self):
        paths = app.openapi()["paths"]
        for path in (
            "/api/v1/stage1/tasks",
            "/api/v1/stage2/audit/tasks",
            "/api/v1/stage2/advice/tasks",
            "/api/v1/stage2/full/tasks",
            "/api/v1/tasks/{task_id}/result",
            "/api/v1/tasks/{task_id}/files/{file_id}",
            "/api/v1/knowledge/cases",
            "/api/v1/knowledge/cases/{case_id}",
            "/api/v1/knowledge/cases/{case_id}/content",
            "/api/v1/knowledge/cases/{case_id}/restore",
            "/api/v1/knowledge/cases/{case_id}/permanent",
            "/api/v1/knowledge/stats",
            "/api/v1/knowledge/regulations",
            "/api/v1/knowledge/regulations/{regulation_id}/permanent",
            "/api/v1/knowledge/regulations/{regulation_id}/generate-rules",
            "/api/v1/knowledge/rules/{rule_id}/test",
            "/api/v1/knowledge/rules/{rule_id}/publish",
            "/api/v1/agent/config",
            "/api/v1/agent/sessions",
            "/api/v1/agent/sessions/{session_id}",
            "/api/v1/agent/sessions/{session_id}/rename",
            "/api/v1/agent/ask",
            "/api/v1/project-archives/projects",
            "/api/v1/project-archives/projects/{project_id}",
            "/api/v1/project-archives/projects/{project_id}/stages",
            "/api/v1/project-archives/stages/{stage_id}/audit",
            "/api/v1/project-archives/stages/{stage_id}/previous-audits",
        ):
            self.assertIn(path, paths)

    def test_flat_stage1_payload_is_built_with_uploaded_file(self):
        with tempfile.TemporaryDirectory() as folder:
            source = Path(folder) / "letter.pdf"
            source.write_bytes(b"pdf")
            value = _prepare_stage1_input(
                source,
                {
                    "project_stage": "规划",
                    "relative_relationship": "单侧",
                    "structure_method": "盾构",
                    "support_components": ["围护桩"],
                },
                "task-1",
            )
        self.assertEqual(value["case_id"], "task-1")
        self.assertEqual(value["source_documents"][0]["path"], str(source.resolve()))
        self.assertEqual(value["pit"]["support_components"], ["围护桩"])

    def test_legacy_disease_value_requires_review_instead_of_failing_schema(self):
        with tempfile.TemporaryDirectory() as folder:
            source = Path(folder) / "letter.pdf"
            source.write_bytes(b"pdf")
            value = _prepare_stage1_input(source, {"disease_severity": "存在病害"}, "task-2")
        self.assertEqual(value["metro_structure"]["disease_severity"], "未知")
        self.assertIn("严重程度未明确", value["review_context"]["manual_notes"])

    def test_single_protection_zone_location_maps_to_consistent_flags(self):
        expected = {
            "特别保护区": (True, True),
            "控制保护区（非特别保护区）": (True, False),
            "保护区外": (False, False),
            "待判断": (None, None),
        }
        with tempfile.TemporaryDirectory() as folder:
            source = Path(folder) / "letter.pdf"
            source.write_bytes(b"pdf")
            for location, flags in expected.items():
                value = _prepare_stage1_input(
                    source,
                    {"protection_zone_location": location},
                    f"task-{location}",
                )
                context = value["review_context"]
                self.assertEqual(context["protection_zone_location"], location)
                self.assertEqual(context["is_in_control_protection_zone"], flags[0])
                self.assertEqual(context["is_in_special_protection_zone"], flags[1])


class AgentServiceTests(unittest.TestCase):
    def test_configuration_hides_full_api_key(self):
        with tempfile.TemporaryDirectory() as folder:
            service = AgentService(Path(folder) / "siliconflow.json")
            status = service.configure("sk-test-1234567890")
            self.assertTrue(status["configured"])
            self.assertEqual(status["provider"], "siliconflow")
            self.assertEqual(status["model"], "Qwen/Qwen3.5-35B-A3B")
            self.assertNotIn("sk-test-1234567890", json.dumps(status))

    def test_general_chat_calls_siliconflow_and_returns_answer(self):
        class FakeResponse:
            def __enter__(self):
                return self

            def __exit__(self, *args):
                return False

            def read(self):
                return json.dumps({
                    "model": "Qwen/Qwen3.5-35B-A3B",
                    "choices": [{"message": {"content": "测试回答"}}],
                    "usage": {"total_tokens": 12},
                }).encode("utf-8")

        with tempfile.TemporaryDirectory() as folder:
            service = AgentService(Path(folder) / "siliconflow.json")
            service.configure("sk-test-1234567890")
            with patch("audit_api.agent.urllib.request.urlopen", return_value=FakeResponse()) as mocked:
                result = service.chat("你好", [], [], False)
            self.assertEqual(result["answer"], "测试回答")
            self.assertEqual(result["provider"], "siliconflow")
            request = mocked.call_args.args[0]
            body = json.loads(request.data.decode("utf-8"))
            self.assertEqual(body["model"], "Qwen/Qwen3.5-35B-A3B")
            self.assertEqual(body["messages"][-1]["content"], "你好")
            self.assertNotIn("知识库材料", body["messages"][0]["content"])


    def test_transient_timeout_is_retried(self):
        class FakeResponse:
            def __enter__(self):
                return self

            def __exit__(self, *args):
                return False

            def read(self):
                return json.dumps({
                    "choices": [{"message": {"content": '{"ok": true}'}}],
                }).encode("utf-8")

        with tempfile.TemporaryDirectory() as folder:
            service = AgentService(Path(folder) / "siliconflow.json")
            service.configure("sk-test-1234567890")
            with (
                patch(
                    "audit_api.agent.urllib.request.urlopen",
                    side_effect=[TimeoutError("slow provider"), FakeResponse()],
                ) as mocked,
                patch("audit_api.agent.time.sleep"),
            ):
                result = service.complete_json("only json", "return ok", max_tokens=20)
        self.assertEqual(result, {"ok": True})
        self.assertEqual(mocked.call_count, 2)

    def test_incomplete_json_is_repaired_locally_without_second_api_call(self):
        class FakeResponse:
            def __enter__(self):
                return self

            def __exit__(self, *args):
                return False

            def read(self):
                return json.dumps({
                    "choices": [{"message": {"content": '{"findings":[{"title":"净距不足"}]'}}],
                }).encode("utf-8")

        with tempfile.TemporaryDirectory() as folder:
            service = AgentService(Path(folder) / "siliconflow.json")
            service.configure("sk-test-1234567890")
            with patch("audit_api.agent.urllib.request.urlopen", return_value=FakeResponse()) as mocked:
                result = service.complete_json("only json", "return findings", max_tokens=20)
        self.assertEqual(result["findings"][0]["title"], "净距不足")
            self.assertEqual(mocked.call_count, 1)


class AgentConversationRepositoryTests(unittest.TestCase):
    def test_session_messages_are_persisted_and_user_scoped(self):
        with tempfile.TemporaryDirectory() as folder:
            repository = AgentConversationRepository(Path(folder) / "agent_chat.sqlite3")
            session = repository.create("user-a", mode="knowledge")
            repository.add_message("user-a", session["session_id"], "user", "查询基坑降水要求", mode="knowledge")
            repository.add_message(
                "user-a",
                session["session_id"],
                "assistant",
                "应核实降水对地铁结构的影响。",
                model="Qwen/Qwen3.5-35B-A3B",
                sources=[{"case_name": "测试案例", "excerpt": "降水控制"}],
                title="基坑降水安全控制",
            )

            loaded = repository.get("user-a", session["session_id"])
            self.assertEqual(loaded["message_count"], 2)
            self.assertEqual(loaded["title"], "基坑降水安全控制")
            self.assertEqual(loaded["messages"][1]["sources"][0]["case_name"], "测试案例")
            self.assertEqual(repository.list("user-b"), [])


class TaskManagerTests(unittest.TestCase):
    def test_successful_task_is_persisted(self):
        with tempfile.TemporaryDirectory() as folder:
            manager = TaskManager(Path(folder), max_workers=1)
            source = Path(folder) / "source.pdf"
            source.write_bytes(b"x")

            def worker(task_id, progress):
                progress("处理中")
                return {"task_id": task_id, "artifact_roots": []}

            task = manager.create("test", source, worker)
            for _ in range(100):
                current = manager.get(task["task_id"])
                if current["status"] in {"success", "failed"}:
                    break
                time.sleep(0.01)
            self.assertEqual(current["status"], "success")
            self.assertEqual(current["progress"], 100)
            self.assertTrue((Path(folder) / f"{task['task_id']}.json").exists())


class KnowledgeBaseTests(unittest.TestCase):
    def test_word_case_is_extracted_exported_and_can_be_disabled(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "测试基坑案例.docx"
            document = Document()
            document.add_paragraph("测试基坑工程位于轨道交通1号线盾构区间单侧，基坑深度8.5m。")
            document.add_paragraph("6.2 审核意见")
            document.add_paragraph("1、施工前应编制专项监测方案，并加强轨道结构变形监测。")
            document.add_paragraph("2、基坑开挖期间应控制地下水位变化，落实应急处置措施。")
            document.add_paragraph("7 结论")
            document.save(source)

            match_database = root / "case_advice_database.json"
            match_database.write_text(
                json.dumps({"format_version": "case_advice_database_v1", "records": []}, ensure_ascii=False),
                encoding="utf-8",
            )
            knowledge = KnowledgeBase(
                database=root / "knowledge.sqlite3",
                file_root=root / "files",
                match_database=match_database,
                python_executable=Path(__import__("sys").executable),
            )
            messages = []
            item = knowledge.import_case(source, "测试基坑案例", "基坑", messages.append)

            self.assertEqual(item["status"], "ready")
            self.assertGreaterEqual(item["advice_count"], 1)
            self.assertGreater(item["text_length"], 0)
            exported = json.loads(match_database.read_text(encoding="utf-8"))
            self.assertEqual(exported["case_count"], 1)
            self.assertEqual(exported["records"][0]["case_id"], item["case_id"])
            self.assertEqual(exported["records"][0]["original_file_name"], "测试基坑案例.docx")

            knowledge.set_active(item["case_id"], False)
            exported = json.loads(match_database.read_text(encoding="utf-8"))
            self.assertEqual(exported["case_count"], 0)
            knowledge.set_active(item["case_id"], True)
            self.assertEqual(knowledge.stats()["matchable"], 1)
            matches = knowledge.search("基坑变形监测", 3)
            self.assertEqual(matches[0]["case_id"], item["case_id"])
            self.assertIn("变形监测", matches[0]["excerpt"])
            self.assertEqual(knowledge.search("测试基坑案例.docx", 3)[0]["case_id"], item["case_id"])
            self.assertEqual(knowledge.search("测试基坑案例", 3)[0]["case_id"], item["case_id"])


class StageTwoAdapterTests(unittest.TestCase):
    def test_audit_adapter_returns_generated_summary(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "plan.docx"
            source.write_bytes(b"docx")
            output = root / "audit"

            def fake_run(command, *, cwd, log_file, progress, label):
                run_dir = output / "run-1"
                run_dir.mkdir(parents=True)
                (run_dir / "audit_summary.json").write_text(
                    json.dumps({"run_dir": str(run_dir), "case_json": str(run_dir / "case.json")}),
                    encoding="utf-8",
                )

            with patch("deepke_case_extract.pipelines._run", side_effect=fake_run):
                result = run_audit_one_plan(source, run_id="run-1", output_root=output)
            self.assertEqual(result["run_dir"], str(output / "run-1"))

    def test_advice_adapter_returns_complete_match(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "case.json"
            source.write_text("{}", encoding="utf-8")
            database = root / "database.json"
            database.write_text("{}", encoding="utf-8")
            output = root / "advice"

            def fake_run(command, *, cwd, log_file, progress, label):
                output.mkdir(parents=True, exist_ok=True)
                (output / "case_advice_match.json").write_text(
                    json.dumps({"format_version": "case_advice_match_v1", "best_match": {"advices": [{"text": "建议原文"}]}}),
                    encoding="utf-8",
                )

            with patch("deepke_case_extract.pipelines._run", side_effect=fake_run):
                result = run_match_new_case_advice(source, output_dir=output, database_file=database)
            self.assertEqual(result["best_match"]["advices"][0]["text"], "建议原文")


class DynamicRegulationRuleTests(unittest.TestCase):
    def sample_rule(self):
        return {
            "name": "上跨隧道竖向净距要求",
            "object": "基坑",
            "relation": "上跨",
            "actual_field": "vertical_clearance",
            "inputs": {
                "vertical_clearance": {"name": "实际竖向净距", "unit": "m", "required": True},
                "tunnel_diameter": {"name": "隧道外径D", "symbol": "D", "unit": "m", "required": True},
            },
            "limit_formula": "max(tunnel_diameter, 4)",
            "operator": ">=",
            "action": {"clearance_compliant": True},
            "source": {"original_text": "5.3.2 上跨隧道，竖向净距不得小于max(1D,4m)。", "context": "5.3.2 上跨隧道，竖向净距不得小于max(1D,4m)。", "clause": "5.3.2", "page": 12},
        }

    def test_safe_formula_and_boundary_values(self):
        self.assertEqual(SafeExpression.evaluate("max(D, 4)", {"D": "6.2"}), SafeExpression._decimal("6.2"))
        rule = self.sample_rule()
        self.assertEqual(RuleEngine.execute(rule, {"vertical_clearance": 6.2, "tunnel_diameter": 6.2})["status"], "matched")
        self.assertEqual(RuleEngine.execute(rule, {"vertical_clearance": 6.1, "tunnel_diameter": 6.2})["status"], "not_matched")
        self.assertEqual(RuleEngine.execute(rule, {"vertical_clearance": "6200mm", "tunnel_diameter": "6.2m"})["status"], "matched")

    def test_missing_field_and_unsafe_formula(self):
        result = RuleEngine.execute(self.sample_rule(), {"vertical_clearance": 8})
        self.assertEqual(result["status"], "insufficient_data")
        with self.assertRaises(ValueError):
            SafeExpression.evaluate("__import__('os').system('dir')", {})

    def test_conditional_and_lookup_rules(self):
        conditional = {
            "rule_type": "conditional_rule", "name": "基坑应监测",
            "conditions": [{"field": "project_type", "operator": "==", "value": "基坑"}],
            "requirement": {"field": "monitoring_required", "operator": "==", "value": True},
            "action": {"monitoring_compliant": True},
            "source": {"original_text": "基坑工程应实施监测。"},
        }
        self.assertEqual(RuleEngine.execute(conditional, {"project_type": "基坑", "monitoring_required": False})["status"], "not_matched")
        trigger = {
            "rule_type": "conditional_rule", "name": "50m内需要评估",
            "conditions": [{"field": "distance", "operator": "<", "value": 50}],
            "action": {"safety_assessment_required": True},
            "source": {"original_text": "距离既有线路小于50m应开展安全评估。"},
        }
        triggered = RuleEngine.execute(trigger, {"distance": 20})
        self.assertEqual(triggered["status"], "matched")
        self.assertTrue(triggered["action"]["safety_assessment_required"])
        lookup = {
            "rule_type": "lookup_table_rule", "name": "影响等级表",
            "selectors": ["project_type", "relative_relationship"], "output_field": "impact_level",
            "rows": [{"when": {"project_type": "基坑", "relative_relationship": "交叉"}, "result": "一级"}],
            "source": {"original_text": "基坑交叉对应一级影响。"},
        }
        result = RuleEngine.execute(lookup, {"project_type": "基坑", "relative_relationship": "交叉"})
        self.assertEqual(result["status"], "derived")
        self.assertEqual(result["derived_value"], "一级")

    def test_trigger_rule_is_not_counted_as_compliant(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "规程.txt"
            source.write_text("距离小于50m时应开展安全评估。", encoding="utf-8")
            repository = RegulationRepository(root / "regulations.sqlite3", root / "files")
            document = repository.import_document(source, "评估规程", "2026", lambda _: None)
            rule = repository.create_rule(document["regulation_id"], {
                "name": "50m范围触发评估", "actual_field": "distance",
                "inputs": {"distance": {"name": "距离", "unit": "m", "required": True}},
                "limit_formula": "50", "operator": "<",
                "action": {"safety_assessment_required": True},
                "source": {"original_text": "距离小于50m时应开展安全评估。", "clause": "3.1.1"},
            })
            repository.publish_rule(rule["rule_id"])
            audit = run_dynamic_regulation_audit({"distance": 20}, repository)
            self.assertEqual(audit["summary"]["triggered"], 1)
            self.assertEqual(audit["summary"]["compliant"], 0)

    def test_complete_table_can_pass_without_identical_second_ai_result(self):
        repository = object.__new__(RegulationRepository)
        source = (
            '<table><tr><td>监测项目</td><td>一级</td><td>二级</td></tr>'
            '<tr><td>竖向位移</td><td>应测</td><td>宜测</td></tr>'
            '<tr><td>结构裂缝</td><td>应测</td><td>可测</td></tr></table>'
        )
        complete = {
            "rule_type": "lookup_table_rule", "name": "监测项目表",
            "selectors": ["impact_level"], "output_field": "monitoring_requirements",
            "rows": [
                {"when": {"impact_level": "一级"}, "result": {"竖向位移": "应测", "结构裂缝": "应测"}},
                {"when": {"impact_level": "二级"}, "result": {"竖向位移": "宜测", "结构裂缝": "可测"}},
            ],
            "source": {"original_text": source, "clause": "7.2.4"},
        }
        review = repository.automatic_review(complete, independent_match=False)
        self.assertTrue(review["auto_publishable"])
        self.assertEqual(review["table_source_coverage"], 1.0)
        incomplete = json.loads(json.dumps(complete, ensure_ascii=False))
        incomplete["rows"][0]["result"].pop("结构裂缝")
        incomplete["rows"][1]["result"].pop("结构裂缝")
        self.assertFalse(repository.automatic_review(incomplete, independent_match=False)["auto_publishable"])

    def test_generic_table_relation_verifier_rejects_swapped_cells(self):
        repository = object.__new__(RegulationRepository)
        source = (
            '<table><tr><td>工程类型</td><td>相对关系</td><td>影响等级</td></tr>'
            '<tr><td>基坑</td><td>交叉</td><td>一级</td></tr>'
            '<tr><td>基坑</td><td>单侧</td><td>二级</td></tr></table>'
        )
        rule = {
            "rule_type": "lookup_table_rule", "name": "影响等级确定表",
            "selectors": ["project_type", "relative_relationship"], "output_field": "impact_level",
            "rows": [
                {"when": {"project_type": "基坑", "relative_relationship": "交叉"}, "result": "一级"},
                {"when": {"project_type": "基坑", "relative_relationship": "单侧"}, "result": "二级"},
            ],
            "source": {"original_text": source, "clause": "5.1.2"},
        }
        review = repository.automatic_review(rule, independent_match=False)
        self.assertTrue(review["auto_publishable"])
        self.assertEqual(review["table_relation_verification"]["matched_relations"], 2)
        swapped = json.loads(json.dumps(rule, ensure_ascii=False))
        swapped["rows"][0]["result"], swapped["rows"][1]["result"] = "二级", "一级"
        self.assertFalse(repository.automatic_review(swapped, independent_match=False)["auto_publishable"])

    def test_docx_import_preserves_table_as_html(self):
        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "table.docx"
            document = Document()
            document.add_paragraph("7.2.4 监测项目应根据下表选择。")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "监测项目"
            table.cell(0, 1).text = "一级"
            table.cell(1, 0).text = "竖向位移"
            table.cell(1, 1).text = "应测"
            document.save(path)
            text, rows, method = extract_regulation(path)
            self.assertEqual(method, "docx_xml")
            self.assertIn("<table>", text)
            self.assertEqual(rows[1]["content_type"], "table")

    def test_dynamic_audit_executes_only_published_active_rules(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "规程.txt"
            source.write_text("基坑工程应实施监测。", encoding="utf-8")
            repository = RegulationRepository(root / "regulations.sqlite3", root / "files")
            document = repository.import_document(source, "监测规程", "2026", lambda _: None)
            rule = repository.create_rule(document["regulation_id"], {
                "rule_type": "conditional_rule", "name": "基坑应监测",
                "conditions": [{"field": "project_type", "operator": "==", "value": "基坑"}],
                "requirement": {"field": "monitoring_required", "operator": "==", "value": True},
                "source": {"original_text": "基坑工程应实施监测。", "clause": "7.1.1"},
            })
            repository.publish_rule(rule["rule_id"])
            audit = run_dynamic_regulation_audit(
                {"project": {"project_type": "基坑"}, "monitoring_required": False}, repository
            )
            self.assertEqual(audit["summary"]["non_compliant"], 1)
            self.assertEqual(audit["results"][0]["regulation_version"], "2026")
            repository.set_document_active(document["regulation_id"], False)
            self.assertEqual(run_dynamic_regulation_audit({}, repository)["published_rule_count"], 0)

    def test_regulation_repository_lifecycle(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "规程.txt"
            source.write_text("5.3.2 上跨隧道，竖向净距不得小于max(1D,4m)。", encoding="utf-8")
            repository = RegulationRepository(root / "regulations.sqlite3", root / "files")
            item = repository.import_document(source, "测试规程", "2026", lambda _: None)
            candidates = repository.candidates(item["regulation_id"])
            self.assertEqual(len(candidates), 1)
            rule = repository.create_rule(item["regulation_id"], self.sample_rule())
            self.assertTrue(rule["validation"]["valid"])
            published = repository.publish_rule(rule["rule_id"])
            self.assertEqual(published["status"], "published")

    def test_clear_consensus_rule_is_auto_published_but_ambiguous_range_is_not(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "规程.txt"
            source.write_text("5.3.2 上跨隧道，竖向净距不得小于max(1D,4m)。\n5.3.3 地下结构距离既有线路外边线50米范围内需要评估。", encoding="utf-8")
            repository = RegulationRepository(root / "regulations.sqlite3", root / "files")
            item = repository.import_document(source, "测试规程", "2026", lambda _: None)
            candidates = repository.candidates(item["regulation_id"])
            lookup = {value["candidate_id"]: value for value in candidates}
            clear = {
                "candidate_id": candidates[0]["candidate_id"], "name": "竖向净距", "object": "基坑", "relation": "上跨",
                "actual_field": "vertical_clearance",
                "inputs": {
                    "vertical_clearance": {"name": "实际竖向净距", "symbol": "", "unit": "m", "required": True},
                    "tunnel_diameter": {"name": "隧道外径D", "symbol": "D", "unit": "m", "required": True},
                },
                "limit_formula": "max(tunnel_diameter, 4)", "operator": ">=", "action": {"clearance_compliant": True},
            }
            created = repository.save_consensus_ai_rules(item["regulation_id"], {"rules": [clear]}, {"rules": [clear]}, lookup)
            self.assertEqual(created[0]["status"], "published")
            self.assertTrue(created[0]["rule"]["automatic_review"]["auto_publishable"])

            ambiguous = {
                "candidate_id": candidates[1]["candidate_id"], "name": "50米范围", "object": "地下结构", "relation": "邻近",
                "actual_field": "distance_to_existing_line",
                "inputs": {"distance_to_existing_line": {"name": "距离", "symbol": "", "unit": "m", "required": True}},
                "limit_formula": "50", "operator": "<=", "action": {"safety_assessment_required": True},
            }
            created = repository.save_consensus_ai_rules(item["regulation_id"], {"rules": [ambiguous]}, {"rules": [ambiguous]}, lookup)
            self.assertEqual(created[0]["status"], "draft")
            self.assertIn("no_ambiguous_range", created[0]["rule"]["automatic_review"]["failed_checks"])

    def test_single_pass_mode_publishes_without_second_ai_or_manual_review(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "规程.txt"
            source.write_text("5.3.2 上跨隧道，竖向净距不得小于max(1D,4m)。", encoding="utf-8")
            repository = RegulationRepository(root / "regulations.sqlite3", root / "files")
            item = repository.import_document(source, "测试规程", "2026", lambda _: None)
            candidate = repository.candidates(item["regulation_id"])[0]
            value = {
                "candidate_id": candidate["candidate_id"], "name": "竖向净距", "object": "隧道", "relation": "上跨",
                "actual_field": "minimum_vertical_clearance",
                "inputs": {
                    "minimum_vertical_clearance": {"name": "竖向净距", "symbol": "", "unit": "m", "required": True},
                    "tunnel_diameter": {"name": "隧道外径D", "symbol": "D", "unit": "m", "required": True},
                },
                "limit_formula": "max(tunnel_diameter,4)", "operator": ">=", "action": {"clearance_compliant": True},
            }
            saved = repository.save_single_pass_ai_rules(
                item["regulation_id"], {"rules": [value]}, {candidate["candidate_id"]: candidate}
            )
            self.assertEqual(saved[0]["status"], "published")
            self.assertEqual(saved[0]["rule"]["automatic_review"]["mode"], "single_pass_llm_v1")
            value["name"] = "竖向净距更新版"
            replaced = repository.save_single_pass_ai_rules(
                item["regulation_id"], {"rules": [value]}, {candidate["candidate_id"]: candidate}
            )
            self.assertEqual(replaced[0]["status"], "published")
            active = repository.list_rules(item["regulation_id"])
            self.assertEqual(len(active), 1)
            self.assertEqual(active[0]["name"], "竖向净距更新版")


if __name__ == "__main__":
    unittest.main()
